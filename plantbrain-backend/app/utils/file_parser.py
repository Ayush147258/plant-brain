"""Utilities for extracting clean text from uploaded PlantBrain files."""

from __future__ import annotations

import logging
import os
import re
import tempfile
import zipfile
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any

from app.config import settings


logger = logging.getLogger(__name__)


class FileParser:
    """Parser helpers for extracting text and metadata from supported file types."""

    IMAGE_FORMATS = {"png", "jpg", "jpeg", "tiff", "bmp"}
    EXTENSION_FORMATS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".txt": "txt",
        ".png": "png",
        ".jpg": "jpg",
        ".jpeg": "jpg",
        ".tif": "tiff",
        ".tiff": "tiff",
        ".bmp": "bmp",
        ".xlsx": "xlsx",
        ".xls": "xls",
        ".dxf": "dxf",
        ".dwg": "dwg",
        ".eml": "eml",
        ".msg": "msg",
    }

    @staticmethod
    def parse_file_sync(file_path: str, file_type: str = "") -> dict[str, Any]:
        """Synchronously route a file to the correct parser, preferring magic-byte detection."""

        detected_format = FileParser.detect_format(file_path)
        normalized_type = FileParser._normalize_requested_type(file_type)
        selected_format = detected_format if detected_format != "unknown" else normalized_type
        logger.info("Parsing file %s as %s", file_path, selected_format)

        parsers = {
            "pdf": FileParser.parse_pdf,
            "docx": FileParser.parse_docx,
            "doc": FileParser.parse_doc,
            "txt": FileParser.parse_txt,
            "png": FileParser.parse_image,
            "jpg": FileParser.parse_image,
            "jpeg": FileParser.parse_image,
            "tiff": FileParser.parse_image,
            "bmp": FileParser.parse_image,
            "xlsx": FileParser.parse_xlsx,
            "xls": FileParser.parse_xls,
            "dxf": FileParser.parse_dxf,
            "dwg": FileParser.parse_dwg,
            "eml": FileParser.parse_eml,
            "msg": FileParser.parse_msg,
        }
        parser = parsers.get(selected_format)
        if parser is None:
            return FileParser._error_result(
                selected_format,
                "unsupported_format",
                f"Unsupported file format: {selected_format or file_type or 'unknown'}",
            )
        return parser(file_path)

    @staticmethod
    async def parse_file(file_path: str, file_type: str) -> dict[str, Any]:
        """Route a file to the correct parser based on detected content."""

        return FileParser.parse_file_sync(file_path, file_type)

    @staticmethod
    def detect_format(path: str) -> str:
        """Detect file format from magic bytes, warning when extension disagrees."""

        file_path = Path(path)
        extension_format = FileParser.EXTENSION_FORMATS.get(file_path.suffix.lower(), "unknown")
        try:
            with open(file_path, "rb") as file:
                head = file.read(8192)
        except OSError as exc:
            logger.warning("Could not read %s for format detection: %s", path, exc)
            return extension_format

        detected = FileParser._detect_format_from_bytes(head, file_path, extension_format)
        if detected != "unknown" and extension_format != "unknown" and detected != extension_format:
            logger.warning(
                "File extension for %s suggests %s but magic bytes indicate %s; using %s",
                path,
                extension_format,
                detected,
                detected,
            )
        return detected if detected != "unknown" else extension_format

    @staticmethod
    def parse_pdf(file_path: str) -> dict[str, Any]:
        """Extract text from a PDF, using confidence-scored OCR for sparse pages."""

        logger.info("Extracting PDF text from %s", file_path)
        if settings.document_parser.lower() == "docling":
            try:
                return FileParser._parse_pdf_docling(file_path)
            except Exception:
                logger.warning("Docling conversion failed; falling back to PyMuPDF/OCR", exc_info=True)
        try:
            import fitz

            page_texts: list[dict[str, Any]] = []
            ocr_confidences: list[float] = []
            with fitz.open(file_path) as doc:
                if doc.needs_pass:
                    return FileParser._error_result("pdf", "password_protected", "PDF is password protected")
                page_count = doc.page_count
                metadata = {
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "pages": page_count,
                }

                for page_index, page in enumerate(doc, start=1):
                    logger.debug("Extracting text from PDF page %s", page_index)
                    text = page.get_text("text") or ""
                    page_confidence = None
                    if len(text.strip()) < 50:
                        logger.info("Page %s has sparse text; attempting OCR", page_index)
                        ocr_result = FileParser._ocr_pdf_page(page)
                        if ocr_result["text"]:
                            text = f"{text}\n{ocr_result['text']}" if text.strip() else ocr_result["text"]
                        page_confidence = ocr_result["confidence"]
                        if page_confidence is not None:
                            ocr_confidences.append(page_confidence)
                    page_texts.append({"page_number": page_index, "text": FileParser.clean_text(text), "ocr_confidence": page_confidence})

            combined_text = FileParser.clean_text("\n\n".join(page["text"] for page in page_texts))
            average_confidence = FileParser._average(ocr_confidences)
            if average_confidence is not None:
                metadata["ocr_average_confidence"] = average_confidence
                metadata["low_confidence"] = average_confidence < settings.ocr_confidence_threshold
            return FileParser._success_result("pdf", combined_text, {**metadata, "parser": "pymupdf"}, page_count, page_texts)
        except Exception as exc:
            logger.exception("Failed to parse PDF %s", file_path)
            return FileParser._error_result("pdf", "parse_error", f"Failed to parse PDF: {exc}")

    @staticmethod
    def _parse_pdf_docling(file_path: str) -> dict[str, Any]:
        """Convert a PDF to structure-preserving Markdown with Docling."""

        from docling.document_converter import DocumentConverter

        result = DocumentConverter().convert(file_path)
        markdown = result.document.export_to_markdown()
        pages = len(getattr(result.document, "pages", {}) or {})
        if not markdown.strip():
            raise ValueError("Docling returned no document content")
        return FileParser._success_result(
            "pdf",
            FileParser.clean_text(markdown),
            {"parser": "docling", "format": "markdown", "pages": pages},
            pages,
        )

    @staticmethod
    def parse_docx(file_path: str) -> dict[str, Any]:
        """Extract paragraph and table text from a Word document."""

        logger.info("Extracting DOCX text from %s", file_path)
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts: list[str] = []
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            text_parts.extend(paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return FileParser._success_result(
                "docx",
                FileParser.clean_text("\n".join(text_parts)),
                {"paragraphs": len(paragraphs), "tables": len(doc.tables)},
                0,
            )
        except Exception as exc:
            logger.exception("Failed to parse DOCX %s", file_path)
            return FileParser._error_result("docx", "parse_error", f"Failed to parse DOCX: {exc}")

    @staticmethod
    def parse_doc(file_path: str) -> dict[str, Any]:
        """Return a clean failure for legacy binary Word files."""

        return FileParser._error_result(
            "doc",
            "unsupported_legacy_doc",
            "Legacy .doc files are detected but not parsed; convert to DOCX or PDF first.",
        )

    @staticmethod
    def parse_txt(file_path: str) -> dict[str, Any]:
        """Read text from a plain text file with UTF-8 and Latin-1 fallback."""

        logger.info("Extracting TXT text from %s", file_path)
        try:
            detected_encoding = "utf-8"
            try:
                with open(file_path, encoding="utf-8") as file:
                    text = file.read()
            except UnicodeDecodeError:
                detected_encoding = "latin-1"
                with open(file_path, encoding="latin-1") as file:
                    text = file.read()
            return FileParser._success_result("txt", FileParser.clean_text(text), {"encoding": detected_encoding}, 1)
        except Exception as exc:
            logger.exception("Failed to parse TXT %s", file_path)
            return FileParser._error_result("txt", "parse_error", f"Failed to parse text file: {exc}")

    @staticmethod
    def parse_image(file_path: str) -> dict[str, Any]:
        """Extract text from an image or multi-page TIFF using preprocessed confidence-scored OCR."""

        detected_format = FileParser.detect_format(file_path)
        image_format = detected_format if detected_format in FileParser.IMAGE_FORMATS else "image"
        logger.info("Extracting image text from %s", file_path)
        try:
            from PIL import Image, ImageSequence

            page_texts: list[dict[str, Any]] = []
            page_count = 0
            confidences: list[float] = []
            with Image.open(file_path) as image:
                metadata = {"width": image.width, "height": image.height, "mode": image.mode}
                for index, frame in enumerate(ImageSequence.Iterator(image), start=1):
                    page_count += 1
                    ocr_result = FileParser.ocr_image(frame.convert("RGB"))
                    if ocr_result["confidence"] is not None:
                        confidences.append(ocr_result["confidence"])
                    page_texts.append(
                        {
                            "page_number": index,
                            "text": FileParser.clean_text(ocr_result["text"]),
                            "ocr_confidence": ocr_result["confidence"],
                        }
                    )
            average_confidence = FileParser._average(confidences)
            low_confidence = average_confidence is not None and average_confidence < settings.ocr_confidence_threshold
            metadata.update({"ocr_average_confidence": average_confidence, "low_confidence": low_confidence})
            combined_text = FileParser.clean_text("\n\n".join(page["text"] for page in page_texts))
            return FileParser._success_result(image_format, combined_text, metadata, page_count, page_texts)
        except Exception as exc:
            logger.exception("Failed to parse image %s", file_path)
            return FileParser._error_result(image_format, "parse_error", f"Failed to parse image: {exc}")

    @staticmethod
    def preprocess_image_for_ocr(image: Any) -> Any:
        """Deskew, denoise, and threshold an image before OCR."""

        import cv2
        import numpy as np
        from PIL import Image

        gray = np.array(image.convert("L"))
        angle = FileParser._estimate_skew_angle(gray)
        if abs(angle) > 0.1:
            height, width = gray.shape[:2]
            matrix = cv2.getRotationMatrix2D((width / 2, height / 2), angle, 1.0)
            gray = cv2.warpAffine(gray, matrix, (width, height), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
        denoised = cv2.fastNlMeansDenoising(gray, None, h=30, templateWindowSize=7, searchWindowSize=21)
        thresholded = cv2.adaptiveThreshold(
            denoised,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            11,
        )
        return Image.fromarray(thresholded)

    @staticmethod
    def ocr_image(image: Any) -> dict[str, Any]:
        """Run OCR with per-word confidence and return text plus average confidence."""

        import pytesseract
        from pytesseract import Output

        processed = FileParser.preprocess_image_for_ocr(image)
        data = pytesseract.image_to_data(processed, lang="eng+hin", output_type=Output.DICT)
        words: list[str] = []
        confidences: list[float] = []
        for word, confidence in zip(data.get("text", []), data.get("conf", []), strict=False):
            clean_word = str(word).strip()
            try:
                confidence_value = float(confidence)
            except (TypeError, ValueError):
                confidence_value = -1.0
            if clean_word:
                words.append(clean_word)
            if confidence_value >= 0:
                confidences.append(confidence_value)
        return {"text": FileParser.clean_text(" ".join(words)), "confidence": FileParser._average(confidences)}

    @staticmethod
    def parse_xlsx(file_path: str) -> dict[str, Any]:
        """Extract cell text from an XLSX workbook."""

        logger.info("Extracting XLSX text from %s", file_path)
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
            text, sheet_count = FileParser._workbook_text_from_openpyxl(workbook)
            workbook.close()
            return FileParser._success_result("xlsx", text, {"sheets": sheet_count, "parser": "openpyxl"}, sheet_count)
        except Exception as exc:
            logger.exception("Failed to parse XLSX %s", file_path)
            return FileParser._error_result("xlsx", "parse_error", f"Failed to parse XLSX: {exc}")

    @staticmethod
    def parse_xls(file_path: str) -> dict[str, Any]:
        """Extract cell text from a legacy XLS workbook."""

        logger.info("Extracting XLS text from %s", file_path)
        try:
            import xlrd

            workbook = xlrd.open_workbook(file_path)
            lines: list[str] = []
            for sheet in workbook.sheets():
                lines.append(f"Sheet: {sheet.name}")
                for row_index in range(sheet.nrows):
                    values = [str(sheet.cell_value(row_index, col)).strip() for col in range(sheet.ncols)]
                    row_text = " | ".join(value for value in values if value)
                    if row_text:
                        lines.append(row_text)
            return FileParser._success_result("xls", FileParser.clean_text("\n".join(lines)), {"sheets": workbook.nsheets, "parser": "xlrd"}, workbook.nsheets)
        except Exception as exc:
            logger.exception("Failed to parse XLS %s", file_path)
            return FileParser._error_result("xls", "parse_error", f"Failed to parse XLS: {exc}")

    @staticmethod
    def parse_eml(file_path: str) -> dict[str, Any]:
        """Extract headers and plain text body from an RFC 822 email file."""

        logger.info("Extracting EML text from %s", file_path)
        try:
            with open(file_path, "rb") as file:
                message = BytesParser(policy=policy.default).parse(file)
            parts = [
                f"From: {message.get('from', '')}",
                f"To: {message.get('to', '')}",
                f"Subject: {message.get('subject', '')}",
                f"Date: {message.get('date', '')}",
            ]
            body_parts: list[str] = []
            if message.is_multipart():
                for part in message.walk():
                    if part.get_content_type() == "text/plain":
                        body_parts.append(str(part.get_content()))
            elif message.get_content_type() == "text/plain":
                body_parts.append(str(message.get_content()))
            parts.extend(body_parts)
            return FileParser._success_result("eml", FileParser.clean_text("\n".join(parts)), {"subject": message.get("subject", "")}, 1)
        except Exception as exc:
            logger.exception("Failed to parse EML %s", file_path)
            return FileParser._error_result("eml", "parse_error", f"Failed to parse EML: {exc}")

    @staticmethod
    def parse_msg(file_path: str) -> dict[str, Any]:
        """Extract text from an Outlook MSG file using extract-msg."""

        logger.info("Extracting MSG text from %s", file_path)
        try:
            import extract_msg

            message = extract_msg.Message(file_path)
            parts = [
                f"From: {message.sender or ''}",
                f"To: {message.to or ''}",
                f"Subject: {message.subject or ''}",
                f"Date: {message.date or ''}",
                message.body or "",
            ]
            subject = message.subject or ""
            try:
                message.close()
            except Exception:
                pass
            return FileParser._success_result("msg", FileParser.clean_text("\n".join(parts)), {"subject": subject}, 1)
        except Exception as exc:
            logger.exception("Failed to parse MSG %s", file_path)
            return FileParser._error_result("msg", "parse_error", f"Failed to parse MSG: {exc}")

    @staticmethod
    def parse_dxf(file_path: str) -> dict[str, Any]:
        """Extract text-like entities and layer metadata from a DXF drawing."""

        logger.info("Extracting DXF text from %s", file_path)
        try:
            import ezdxf

            document = ezdxf.readfile(file_path)
            lines: list[str] = []
            for entity in document.modelspace():
                entity_type = entity.dxftype()
                if entity_type in {"TEXT", "MTEXT"}:
                    value = entity.plain_text() if hasattr(entity, "plain_text") else entity.dxf.text
                    lines.append(str(value))
                elif entity_type in {"INSERT", "LINE", "LWPOLYLINE", "POLYLINE", "CIRCLE", "ARC"}:
                    layer = getattr(entity.dxf, "layer", "")
                    lines.append(f"{entity_type} layer={layer}")
            return FileParser._success_result("dxf", FileParser.clean_text("\n".join(lines)), {"parser": "ezdxf"}, 1)
        except Exception as exc:
            logger.exception("Failed to parse DXF %s", file_path)
            return FileParser._error_result("dxf", "parse_error", f"Failed to parse DXF: {exc}")

    @staticmethod
    def parse_dwg(file_path: str) -> dict[str, Any]:
        """Return a clean unsupported result for binary DWG drawings."""

        return FileParser._error_result(
            "dwg",
            "unsupported_dwg",
            "DWG binary parsing is unsupported. Convert the drawing to DXF or PDF first.",
        )

    @staticmethod
    def clean_text(text: str) -> str:
        """Normalize extracted text by removing null bytes and excessive whitespace."""

        if not text:
            return ""
        cleaned = text.replace("\x00", "")
        cleaned = "\n".join(line.strip() for line in cleaned.splitlines())
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
        return cleaned.strip()

    @staticmethod
    def detect_file_type(filename: str) -> str:
        """Detect the supported parser type from a filename extension for upload validation."""

        extension_format = FileParser.EXTENSION_FORMATS.get(Path(filename).suffix.lower(), "unknown")
        if extension_format in FileParser.IMAGE_FORMATS:
            return "image"
        if extension_format == "doc":
            return "docx"
        return extension_format

    @staticmethod
    def _detect_format_from_bytes(head: bytes, file_path: Path, extension_format: str) -> str:
        if head.startswith(b"%PDF"):
            return "pdf"
        if head.startswith(b"\x89PNG\r\n\x1a\n"):
            return "png"
        if head.startswith(b"\xff\xd8\xff"):
            return "jpg"
        if head.startswith((b"II*\x00", b"MM\x00*")):
            return "tiff"
        if head.startswith(b"BM"):
            return "bmp"
        if head.startswith(b"AC10"):
            return "dwg"
        if head.startswith(b"PK\x03\x04") or head.startswith(b"PK\x05\x06") or head.startswith(b"PK\x07\x08"):
            return FileParser._detect_zip_format(file_path, extension_format)
        if head.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
            return extension_format if extension_format in {"doc", "xls", "msg"} else "doc"
        if FileParser._looks_like_dxf(head):
            return "dxf"
        if FileParser._looks_like_eml(head):
            return "eml"
        if FileParser._looks_like_text(head):
            return "txt"
        return "unknown"

    @staticmethod
    def _detect_zip_format(file_path: Path, extension_format: str) -> str:
        try:
            with zipfile.ZipFile(file_path) as archive:
                names = set(archive.namelist())
            if "word/document.xml" in names:
                return "docx"
            if "xl/workbook.xml" in names:
                return "xlsx"
        except Exception:
            return extension_format if extension_format in {"docx", "xlsx"} else "unknown"
        return extension_format if extension_format in {"docx", "xlsx"} else "unknown"

    @staticmethod
    def _looks_like_text(head: bytes) -> bool:
        if not head or b"\x00" in head:
            return False
        try:
            head.decode("utf-8")
            return True
        except UnicodeDecodeError:
            try:
                head.decode("latin-1")
                return True
            except UnicodeDecodeError:
                return False

    @staticmethod
    def _looks_like_dxf(head: bytes) -> bool:
        sample = head[:1024].decode("latin-1", errors="ignore").upper()
        return "SECTION" in sample and ("HEADER" in sample or "ENTITIES" in sample or "EOF" in sample)

    @staticmethod
    def _looks_like_eml(head: bytes) -> bool:
        sample = head[:4096].decode("latin-1", errors="ignore").lower()
        return ("subject:" in sample or "from:" in sample) and ("mime-version:" in sample or "content-type:" in sample)

    @staticmethod
    def _normalize_requested_type(file_type: str) -> str:
        normalized = file_type.lower().strip()
        if normalized == "image":
            return "png"
        if normalized == "docx":
            return "docx"
        return normalized or "unknown"

    @staticmethod
    def _ocr_pdf_page(page: Any) -> dict[str, Any]:
        """Render a PDF page to a temporary image and OCR it."""

        temp_path = ""
        try:
            from PIL import Image

            pix = page.get_pixmap(dpi=200)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
                temp_path = temp_file.name
                pix.save(temp_path)
            with Image.open(temp_path) as image:
                return FileParser.ocr_image(image)
        except Exception:
            logger.exception("Failed to OCR sparse PDF page")
            return {"text": "", "confidence": None}
        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)

    @staticmethod
    def _estimate_skew_angle(gray_image: Any) -> float:
        """Estimate the text skew angle in degrees for a grayscale image array."""

        import cv2
        import numpy as np

        coords = np.column_stack(np.where(gray_image < 240))
        if len(coords) < 10:
            return 0.0
        angle = cv2.minAreaRect(coords)[-1]
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle
        return float(angle)

    @staticmethod
    def _workbook_text_from_openpyxl(workbook: Any) -> tuple[str, int]:
        lines: list[str] = []
        for sheet in workbook.worksheets:
            lines.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(value).strip() for value in row if value is not None and str(value).strip())
                if row_text:
                    lines.append(row_text)
        return FileParser.clean_text("\n".join(lines)), len(workbook.worksheets)

    @staticmethod
    def _average(values: list[float]) -> float | None:
        if not values:
            return None
        return round(sum(values) / len(values), 2)

    @staticmethod
    def _success_result(
        detected_format: str,
        text: str,
        metadata: dict[str, Any],
        pages: int,
        page_texts: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        result = {
            "text": text,
            "metadata": {**metadata, "format": detected_format},
            "pages": pages,
            "success": True,
            "format": detected_format,
            "error": None,
            "ocr_average_confidence": metadata.get("ocr_average_confidence"),
            "low_confidence": bool(metadata.get("low_confidence", False)),
        }
        if page_texts is not None:
            result["page_texts"] = page_texts
        return result

    @staticmethod
    def _error_result(detected_format: str, reason: str, error: str) -> dict[str, Any]:
        """Build a standard parser failure response."""

        return {
            "text": "",
            "metadata": {},
            "pages": 0,
            "success": False,
            "format": detected_format,
            "reason": reason,
            "error": error,
        }
