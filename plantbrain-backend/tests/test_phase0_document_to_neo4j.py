"""Phase 0 document-to-Neo4j integration coverage.

These tests require a disposable Neo4j database. They intentionally do not use the
production NEO4J_* variables unless PLANTBRAIN_TEST_NEO4J_URI is set, so local test
runs cannot accidentally mutate a live graph.
"""

from __future__ import annotations

import os
from pathlib import Path
from uuid import uuid4

import fitz
import pytest
from docx import Document
from PIL import Image, ImageDraw

from app.config import settings
from app.services.multimodal_extraction_service import multimodal_extraction_service
from app.services.neo4j_service import neo4j_service
from app.utils.file_parser import FileParser


def _test_neo4j_config() -> tuple[str, str, str]:
    uri = os.getenv("PLANTBRAIN_TEST_NEO4J_URI", "")
    user = os.getenv("PLANTBRAIN_TEST_NEO4J_USER", "neo4j")
    password = os.getenv("PLANTBRAIN_TEST_NEO4J_PASSWORD", "")
    if not uri or not password:
        pytest.skip("Set PLANTBRAIN_TEST_NEO4J_URI/USER/PASSWORD for disposable Neo4j integration tests")
    return uri, user, password


@pytest.fixture()
def neo4j_test_driver(monkeypatch):
    uri, user, password = _test_neo4j_config()
    monkeypatch.setattr(settings, "neo4j_uri", uri)
    monkeypatch.setattr(settings, "neo4j_user", user)
    monkeypatch.setattr(settings, "neo4j_password", password)

    from neo4j import GraphDatabase

    driver = GraphDatabase.driver(uri, auth=(user, password))
    try:
        with driver.session() as session:
            session.run("RETURN 1").consume()
        yield driver
    finally:
        driver.close()


@pytest.fixture(autouse=True)
def deterministic_multimodal(monkeypatch):
    def fake_generate_json(file_path: str, prompt: str, schema: dict):
        name = Path(file_path).name.lower()
        if "pid" in name:
            return {
                "zone": "PHASE0-ZONE",
                "equipment": [
                    {"id": "P-901", "type": "pump", "confidence": "high"},
                    {"id": "HX-901", "type": "heat_exchanger", "confidence": "high"},
                ],
                "valves": [
                    {
                        "valve_id": "V-901",
                        "valve_type": "gate",
                        "connects_from": "P-901",
                        "connects_to": "HX-901",
                        "confidence": "high",
                    }
                ],
                "instruments": [
                    {"tag": "PT-901", "attached_to_line_between": ["P-901", "HX-901"], "confidence": "high"}
                ],
                "confidence_flags": [],
            }
        return {
            "entries": [
                {
                    "Asset_ID": "P-902",
                    "Failure_Mode": "seal leakage",
                    "Date": "2026-07-07",
                    "Technician_Notes": "Seal leakage observed during inspection.",
                    "confidence": "high",
                }
            ],
            "confidence_flags": [],
        }

    monkeypatch.setattr(multimodal_extraction_service, "_generate_json", fake_generate_json)


@pytest.fixture()
def source_id() -> str:
    return f"phase0-{uuid4()}"


def _cleanup_source(driver, source_document_id: str) -> None:
    query = """
    MATCH (n)
    WHERE n.source_document_id = $source_document_id
       OR n.id IN ['P-901', 'HX-901', 'V-901', 'PT-901', 'P-902']
       OR n.name = 'PHASE0-ZONE'
    DETACH DELETE n
    """
    with driver.session() as session:
        session.run(query, source_document_id=source_document_id).consume()


def _assert_pid_graph(driver) -> None:
    query = """
    MATCH (p:Equipment {id: 'P-901'})-[r:CONNECTED_THROUGH {valve_id: 'V-901'}]->(hx:Equipment {id: 'HX-901'})
    MATCH (v:Valve {id: 'V-901'})
    MATCH (i:Instrument {id: 'PT-901'})
    MATCH (:Zone {name: 'PHASE0-ZONE'})-[:CONTAINS]->(p)
    RETURN count(*) AS count
    """
    with driver.session() as session:
        assert session.run(query).single()["count"] == 1


def _assert_log_graph(driver) -> None:
    query = """
    MATCH (:Equipment {id: 'P-902'})-[:HAD_EVENT]->(m:MaintenanceEvent {failure_mode: 'seal leakage'})
    RETURN count(m) AS count
    """
    with driver.session() as session:
        assert session.run(query).single()["count"] == 1


def _write_txt(path: Path) -> None:
    path.write_text("Maintenance log: P-902 seal leakage observed during inspection.", encoding="utf-8")


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Maintenance record for P-902: seal leakage observed during inspection.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Asset"
    table.cell(0, 1).text = "P-902"
    document.save(path)


def _write_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "P&ID Zone PHASE0-ZONE: P-901 connects through V-901 to HX-901 with PT-901.")
    document.save(path)
    document.close()


def _write_image(path: Path) -> None:
    image = Image.new("RGB", (1200, 320), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 120), "Maintenance log P-902 seal leakage inspection", fill="black")
    image.save(path)


@pytest.mark.integration
@pytest.mark.parametrize(
    ("filename", "file_type", "writer", "kind"),
    [
        ("phase0_log.txt", "txt", _write_txt, "maintenance_log"),
        ("phase0_log.docx", "docx", _write_docx, "maintenance_log"),
        ("phase0_pid.pdf", "pdf", _write_pdf, "pid"),
        ("phase0_log.png", "image", _write_image, "maintenance_log"),
    ],
)
def test_existing_document_formats_parse_extract_merge_and_query_neo4j(
    neo4j_test_driver,
    tmp_path,
    monkeypatch,
    source_id,
    filename,
    file_type,
    writer,
    kind,
) -> None:
    monkeypatch.setattr(settings, "document_parser", "pymupdf")
    _cleanup_source(neo4j_test_driver, source_id)
    file_path = tmp_path / filename
    writer(file_path)

    parse_result = FileParser.parse_file_sync(str(file_path), file_type)
    try:
        assert parse_result["success"] is True
        assert parse_result["text"].strip()

        extracted = multimodal_extraction_service.extract(str(file_path), kind, "PHASE0-ZONE")
        if kind == "pid":
            neo4j_service.merge_pid_extraction(extracted, source_id)
            _assert_pid_graph(neo4j_test_driver)
        else:
            neo4j_service.merge_text_equipment(["P-902"], source_id)
            neo4j_service.merge_log_extraction(extracted, source_id)
            _assert_log_graph(neo4j_test_driver)
    finally:
        _cleanup_source(neo4j_test_driver, source_id)